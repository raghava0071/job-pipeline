# =============================================================================
# COVER_LETTER.PY — Personalized Cover Letter Generator
# Creates a polished .docx cover letter tailored to each job
# =============================================================================

import os
import re
import sys
import pandas as pd
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ python-docx not installed. Run:  pip install python-docx")
    sys.exit(1)

from raghav_profile import PROFILE, SKILLS

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "output", "cover_letters")

COLOR_ACCENT = RGBColor(0x1F, 0x5C, 0x99)
COLOR_DARK   = RGBColor(0x22, 0x22, 0x22)
COLOR_MID    = RGBColor(0x55, 0x55, 0x55)
FONT_NAME    = "Calibri"


def _set_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.name   = FONT_NAME
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _para_space(para, before_pt=0, after_pt=4):
    pPr = para._p.get_or_add_pPr()
    spg = OxmlElement("w:spacing")
    spg.set(qn("w:before"), str(int(before_pt * 20)))
    spg.set(qn("w:after"),  str(int(after_pt  * 20)))
    pPr.append(spg)


def _set_margins(doc, top=1.0, bottom=1.0, left=1.1, right=1.1):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)


def _pick_top_skills(jd_keywords: list[str], n=6) -> list[str]:
    """Return top matching skills from Raghav's stack that appear in JD."""
    kw_set    = set(k.lower() for k in jd_keywords)
    all_skills = [s for group in SKILLS.values() for s in group]
    matched    = [s for s in all_skills if s.lower() in kw_set]
    # Fill with key skills if not enough matches
    fallback   = ["Python", "Apache Spark", "Azure Data Factory", "SQL", "ETL/ELT Pipelines", "PySpark"]
    combined   = list(dict.fromkeys(matched + fallback))
    return combined[:n]


def build_cover_letter(
    job_title:   str,
    company:     str,
    jd_keywords: list[str],
    ats_score:   float,
    jd_preview:  str = "",
    output_path: str = None,
) -> str:
    """Build and save a cover letter .docx. Returns saved file path."""

    doc = Document()
    _set_margins(doc)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)

    today = datetime.now().strftime("%B %d, %Y")
    top_skills = _pick_top_skills(jd_keywords)

    # ── HEADER ────────────────────────────────────────────────────────────────
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_space(p_name, before_pt=0, after_pt=2)
    r = p_name.add_run(PROFILE["name"])
    _set_font(r, 16, bold=True, color=COLOR_ACCENT)

    p_contact = doc.add_paragraph()
    _para_space(p_contact, before_pt=0, after_pt=2)
    rc = p_contact.add_run(
        f"{PROFILE['email']}  ·  {PROFILE['phone']}  ·  {PROFILE['location']}\n"
        f"{PROFILE['linkedin']}  ·  {PROFILE['github']}"
    )
    _set_font(rc, 10, color=COLOR_MID)

    # Divider
    p_div = doc.add_paragraph()
    _para_space(p_div, before_pt=4, after_pt=6)
    pPr  = p_div._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "1F5C99")
    pBdr.append(bot); pPr.append(pBdr)

    # Date
    p_date = doc.add_paragraph()
    _para_space(p_date, before_pt=0, after_pt=4)
    _set_font(p_date.add_run(today), 11, color=COLOR_DARK)

    # Salutation
    p_sal = doc.add_paragraph()
    _para_space(p_sal, before_pt=0, after_pt=4)
    _set_font(p_sal.add_run(f"Dear {company} Hiring Team,"), 11, color=COLOR_DARK)

    # ── BODY ──────────────────────────────────────────────────────────────────
    # Opening paragraph
    opening = (
        f"I am writing to express my strong interest in the {job_title} role at {company}. "
        f"With a Master of Science in Data Science and Analytics from Florida Atlantic University "
        f"and hands-on experience building enterprise-grade data pipelines at Knowvia Tech Inc, "
        f"I am confident in my ability to contribute meaningfully to your team from day one."
    )
    p1 = doc.add_paragraph()
    _para_space(p1, before_pt=0, after_pt=6)
    _set_font(p1.add_run(opening), 11, color=COLOR_DARK)
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Skills paragraph
    skills_str = ", ".join(top_skills[:-1]) + f", and {top_skills[-1]}" if len(top_skills) > 1 else top_skills[0] if top_skills else "Python and SQL"
    body = (
        f"In my current role, I architect and deploy end-to-end ETL/ELT pipelines using "
        f"{skills_str}, processing data at scale across multi-cloud environments including "
        f"Azure, AWS, and GCP. I have built automated data quality frameworks, implemented "
        f"real-time and batch processing systems, and collaborated cross-functionally to "
        f"translate complex business requirements into scalable data architecture. My work "
        f"directly supports analytics, reporting, and machine learning workflows that drive "
        f"business decisions."
    )
    p2 = doc.add_paragraph()
    _para_space(p2, before_pt=0, after_pt=6)
    _set_font(p2.add_run(body), 11, color=COLOR_DARK)
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Why this company
    why = (
        f"I am particularly drawn to {company} because of the opportunity to work on "
        f"meaningful data challenges at scale. I thrive in environments where data engineering "
        f"directly impacts product decisions and business outcomes. I am authorized to work in "
        f"the United States under F-1 OPT/STEM OPT, require no sponsorship, and am available "
        f"to start immediately."
    )
    p3 = doc.add_paragraph()
    _para_space(p3, before_pt=0, after_pt=6)
    _set_font(p3.add_run(why), 11, color=COLOR_DARK)
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Closing
    closing = (
        f"I would welcome the opportunity to discuss how my skills and experience align with "
        f"your team's needs. Thank you for your time and consideration — I look forward to "
        f"connecting."
    )
    p4 = doc.add_paragraph()
    _para_space(p4, before_pt=0, after_pt=10)
    _set_font(p4.add_run(closing), 11, color=COLOR_DARK)
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Sign-off
    p_sign = doc.add_paragraph()
    _para_space(p_sign, before_pt=0, after_pt=2)
    _set_font(p_sign.add_run("Warm regards,"), 11, color=COLOR_DARK)

    p_sig = doc.add_paragraph()
    _para_space(p_sig, before_pt=0, after_pt=2)
    _set_font(p_sig.add_run(PROFILE["name"]), 11, bold=True, color=COLOR_ACCENT)

    p_sig2 = doc.add_paragraph()
    _para_space(p_sig2, before_pt=0, after_pt=2)
    _set_font(p_sig2.add_run(
        f"{PROFILE['email']}  ·  {PROFILE['phone']}"
    ), 10, color=COLOR_MID)

    # ── SAVE ──────────────────────────────────────────────────────────────────
    if not output_path:
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        safe_company = re.sub(r"[^\w\s-]", "", company).strip().replace(" ", "_")[:30]
        safe_title   = re.sub(r"[^\w\s-]", "", job_title).strip().replace(" ", "_")[:25]
        filename     = f"CoverLetter_{safe_company}_{safe_title}.docx"
        output_path  = os.path.join(OUTPUT_DIR, filename)

    doc.save(output_path)
    print(f"  📝 Cover letter saved → {os.path.basename(output_path)}")
    return output_path


def build_all_cover_letters(filtered_csv: str = None) -> list[str]:
    print("\n" + "═" * 60)
    print("  JOB PIPELINE — Step 4: Building Cover Letters")
    print("═" * 60)

    if not filtered_csv:
        filtered_csv = os.path.join(os.path.dirname(__file__), "data", "filtered_jobs.csv")

    if not os.path.exists(filtered_csv):
        print("❌ filtered_jobs.csv not found. Run jd_parser.py first.")
        sys.exit(1)

    df    = pd.read_csv(filtered_csv)
    paths = []

    for _, row in df.iterrows():
        keywords = [k.strip() for k in str(row.get("jd_keywords", "")).split(",") if k.strip()]
        path = build_cover_letter(
            job_title   = str(row.get("title",      "Data Engineer")),
            company     = str(row.get("company",    "Company")),
            jd_keywords = keywords,
            ats_score   = float(row.get("ats_score", 0)),
            jd_preview  = str(row.get("jd_preview", "")),
        )
        paths.append(path)

    print(f"\n✅ {len(paths)} cover letters built → output/cover_letters/")
    print("═" * 60 + "\n")
    return paths


def save_cover_letter(text: str, job_title: str, company: str, output_path: str = None) -> str:
    """
    Save a Claude-generated plain-text cover letter into a formatted .docx.
    Called by master_run.py after ce.write_cover_letter() returns the text.
    Returns the saved file path.
    """
    doc = Document()
    _set_margins(doc)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)

    today = datetime.now().strftime("%B %d, %Y")

    # Header
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_space(p_name, before_pt=0, after_pt=2)
    _set_font(p_name.add_run(PROFILE["name"]), 16, bold=True, color=COLOR_ACCENT)

    p_contact = doc.add_paragraph()
    _para_space(p_contact, before_pt=0, after_pt=2)
    _set_font(p_contact.add_run(
        f"{PROFILE['email']}  ·  {PROFILE['phone']}  ·  {PROFILE['location']}\n"
        f"{PROFILE['linkedin']}  ·  {PROFILE['github']}"
    ), 10, color=COLOR_MID)

    # Divider line
    p_div = doc.add_paragraph()
    _para_space(p_div, before_pt=4, after_pt=6)
    pPr = p_div._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "1F5C99")
    pBdr.append(bot); pPr.append(pBdr)

    # Date
    p_date = doc.add_paragraph()
    _para_space(p_date, before_pt=0, after_pt=8)
    _set_font(p_date.add_run(today), 11, color=COLOR_DARK)

    # Body — split Claude's text into paragraphs and write each
    body_text = text.strip()
    paragraphs = [p.strip() for p in body_text.split("\n\n") if p.strip()]

    for para_text in paragraphs:
        p = doc.add_paragraph()
        _para_space(p, before_pt=0, after_pt=7)
        _set_font(p.add_run(para_text), 11, color=COLOR_DARK)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Sign-off (if not already in text)
    if "warm regards" not in body_text.lower() and "sincerely" not in body_text.lower():
        p_sign = doc.add_paragraph()
        _para_space(p_sign, before_pt=8, after_pt=2)
        _set_font(p_sign.add_run("Warm regards,"), 11, color=COLOR_DARK)

        p_sig = doc.add_paragraph()
        _para_space(p_sig, before_pt=0, after_pt=2)
        _set_font(p_sig.add_run(PROFILE["name"]), 11, bold=True, color=COLOR_ACCENT)

    # Save
    if not output_path:
        cl_dir = os.path.join(os.path.dirname(__file__), "cover_letters")
        os.makedirs(cl_dir, exist_ok=True)
        safe_co = re.sub(r"[^\w\s-]", "", company).strip().replace(" ", "_")[:30]
        safe_ti = re.sub(r"[^\w\s-]", "", job_title).strip().replace(" ", "_")[:25]
        output_path = os.path.join(cl_dir, f"CoverLetter_{safe_co}_{safe_ti}.docx")

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    build_all_cover_letters()
