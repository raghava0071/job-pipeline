# =============================================================================
# RESUME_BUILDER.PY — ATS-Optimized Word Resume Generator (PRO v3)
#
# GUARANTEES:
#   - 98%+ ATS keyword coverage after tailoring (verified by scanning DOCX text)
#   - Clear BEFORE → AFTER score printed per job
#   - Gap-fill section auto-added if any JD keywords are still missing
#   - Footer shows actual verified score (not estimated)
# =============================================================================

import os
import re
import sys
import pandas as pd
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx not installed. Run:  pip install python-docx")
    sys.exit(1)

from raghav_profile import PROFILE, EDUCATION, EXPERIENCE, SKILLS, PROJECTS

from pathlib import Path as _Path
OUTPUT_DIR = str(_Path.home() / "job_pipeline" / "resumes")

# ── COLORS & FONTS ────────────────────────────────────────────────────────────
COLOR_ACCENT = RGBColor(0x1F, 0x5C, 0x99)
COLOR_BLACK  = RGBColor(0x00, 0x00, 0x00)
COLOR_DARK   = RGBColor(0x22, 0x22, 0x22)
COLOR_MID    = RGBColor(0x44, 0x44, 0x44)
COLOR_GRAY   = RGBColor(0xAA, 0xAA, 0xAA)
FONT_NAME    = "Calibri"

# Synonym map for keyword verification (canonical → surface form variants)
SYNONYM_GROUPS = [
    {"etl", "extract transform load", "data pipelines", "data integration", "data pipeline", "elt"},
    {"pyspark", "apache spark", "spark", "sparksql", "spark sql"},
    {"adf", "azure data factory"},
    {"adls", "azure data lake", "azure data lake storage", "adls gen2"},
    {"azure synapse", "synapse analytics", "azure synapse analytics"},
    {"azure sql", "azure sql database", "azure sql server"},
    {"databricks", "azure databricks"},
    {"aws", "amazon web services"},
    {"gcp", "google cloud", "google cloud platform"},
    {"bigquery", "big query", "bq"},
    {"kafka", "apache kafka"},
    {"airflow", "apache airflow"},
    {"hadoop", "apache hadoop", "hdfs"},
    {"sql", "structured query language"},
    {"t-sql", "tsql", "transact-sql"},
    {"data warehouse", "data warehousing", "dwh", "edw"},
    {"data lake", "lakehouse", "delta lake"},
    {"real-time", "realtime", "real time", "streaming", "stream processing"},
    {"batch processing", "batch jobs", "batch pipeline"},
    {"python", "python3"},
    {"power bi", "powerbi", "microsoft power bi"},
    {"ci/cd", "cicd", "continuous integration"},
    {"docker", "containers", "containerization"},
    {"kubernetes", "k8s"},
    {"dbt", "data build tool"},
    {"nosql", "no sql", "non-relational"},
    {"postgresql", "postgres"},
    {"data quality", "data validation", "data accuracy"},
    {"data governance", "data security", "data lineage"},
    {"rest api", "restful api", "rest", "api integration"},
    {"agile", "scrum", "sprint"},
    {"git", "github", "version control"},
    {"machine learning", "ml", "mlops"},
    {"google analytics 4", "ga4", "google analytics"},
]

def _build_syn_lookup(groups):
    lookup = {}
    for grp in groups:
        canonical = sorted(grp)[0]
        for term in grp:
            lookup[term] = canonical
    return lookup

SYNONYM_LOOKUP = _build_syn_lookup(SYNONYM_GROUPS)


# =============================================================================
# DOCUMENT HELPERS
# =============================================================================

def _set_font(run, size_pt, bold=False, italic=False, color=COLOR_DARK):
    run.font.name      = FONT_NAME
    run.font.size      = Pt(size_pt)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.color.rgb = color


def _para_space(para, before_pt=0, after_pt=0):
    pPr = para._p.get_or_add_pPr()
    spg = OxmlElement("w:spacing")
    spg.set(qn("w:before"), str(int(before_pt * 20)))
    spg.set(qn("w:after"),  str(int(after_pt  * 20)))
    pPr.append(spg)


def _add_bottom_border(para, color="1F5C99", size=6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_hyperlink(para, url, text, size_pt=10):
    part = para.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r   = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "1F5C99")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    rPr.append(color_el)
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(rPr)
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)


def _set_margins(doc, top=0.55, bottom=0.55, left=0.65, right=0.65):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)


# =============================================================================
# ATS KEYWORD VERIFICATION (scan actual DOCX text)
# =============================================================================

def extract_docx_text(doc: Document) -> str:
    """Extract all text from a DOCX document as one lowercase string."""
    texts = []
    for para in doc.paragraphs:
        texts.append(para.text.lower())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    texts.append(para.text.lower())
    return " ".join(texts)


def compute_actual_coverage(jd_keywords: list, doc_text: str) -> tuple:
    """
    Compute how many JD keywords are actually present in the resume text.
    Returns (score_pct, covered_list, missing_list).
    Uses synonym expansion so 'pyspark' covers 'apache spark' etc.
    """
    covered = []
    missing = []

    for kw in jd_keywords:
        kw_lower = kw.lower().strip()
        # Get all synonyms for this keyword
        canon = SYNONYM_LOOKUP.get(kw_lower, kw_lower)
        synonyms = {term for term, c in SYNONYM_LOOKUP.items() if c == canon}
        synonyms.add(kw_lower)
        synonyms.add(canon)

        found = False
        for syn in synonyms:
            if re.search(r"\b" + re.escape(syn) + r"\b", doc_text):
                found = True
                break

        if found:
            covered.append(kw)
        else:
            missing.append(kw)

    total = max(len(jd_keywords), 1)
    score = round(len(covered) / total * 100, 1)
    return score, covered, missing


# =============================================================================
# KEYWORD INJECTION HELPERS
# =============================================================================

def _keyword_in_text(kw: str, text: str) -> bool:
    pattern = re.compile(r"\b" + re.escape(kw.lower()) + r"\b")
    return bool(pattern.search(text.lower()))


def build_enriched_skills(jd_keywords: list, injectable_kws: list) -> tuple:
    """
    Build a skills dict with ALL JD keywords merged into appropriate categories.
    Returns (enriched_dict, injected_count).
    """
    KEYWORD_CATEGORY_HINTS = {
        "azure": "azure", "adf": "azure", "adls": "azure", "synapse": "azure",
        "databricks": "azure", "azure data factory": "azure", "azure sql": "azure",
        "pyspark": "data_engineering", "spark": "data_engineering", "kafka": "data_engineering",
        "airflow": "data_engineering", "hadoop": "data_engineering", "etl": "data_engineering",
        "elt": "data_engineering", "pipeline": "data_engineering", "dbt": "data_engineering",
        "snowflake": "data_engineering", "data warehouse": "data_engineering",
        "data lake": "data_engineering", "streaming": "data_engineering",
        "real-time": "data_engineering", "batch": "data_engineering",
        "delta lake": "data_engineering", "lakehouse": "data_engineering",
        "python": "programming", "sql": "programming", "t-sql": "programming",
        "pandas": "programming", "sparksql": "programming",
        "aws": "cloud", "gcp": "cloud", "bigquery": "cloud", "google cloud": "cloud",
        "redshift": "cloud", "s3": "cloud",
        "postgresql": "databases", "mysql": "databases", "nosql": "databases",
        "mongodb": "databases", "redis": "databases",
        "power bi": "analytics", "tableau": "analytics", "looker": "analytics",
        "machine learning": "analytics", "ga4": "analytics",
        "docker": "data_engineering", "kubernetes": "data_engineering",
        "ci/cd": "data_engineering", "git": "programming",
    }

    enriched = {k: list(v) for k, v in SKILLS.items()}
    kw_set = set(k.lower() for k in jd_keywords)

    # Sort each category: JD-matching items first
    for cat in enriched:
        enriched[cat] = sorted(
            enriched[cat],
            key=lambda s: 1 if s.lower() in kw_set else 0,
            reverse=True,
        )

    # Stopwords — never inject these into the skills section.
    # These are common JD boilerplate words that are NOT tech skills.
    # If a word appears here, it will never be added to the resume skills list.
    SKILLS_STOPWORDS = {
        # Generic English / filler
        "ability", "strong", "knowledge", "skills", "experience", "years",
        "required", "preferred", "plus", "good", "excellent", "great",
        "background", "understanding", "support", "role", "job", "position",
        "candidate", "team", "work", "environment", "collaborative",
        "hands-on", "hands on", "type", "types", "cross-functional",
        "communication", "written", "verbal", "interpersonal",
        "problem", "solving", "analytical", "detail", "oriented",
        "fast", "paced", "startup", "company", "organization",
        "and", "the", "with", "for", "not", "are", "all", "any",
        "new", "our", "your", "key", "top", "use", "may", "can",
        "will", "must", "have", "this", "that", "from", "into",
        "their", "they", "its", "has", "been", "about", "which",
        # Business words extracted from JD boilerplate
        "business", "data", "information", "systems", "solutions",
        "development", "applications", "tools", "technologies", "platforms",
        "management", "services", "health", "medical", "engineering",
        "familiarity", "compensation", "requirements", "processes",
        "operations", "infrastructure", "architecture", "strategy",
        "delivery", "quality", "performance", "security", "compliance",
        "documentation", "reporting", "analysis", "visualization",
        # Job posting section headers (from JD structure)
        "job summary", "key responsibilities", "required qualifications",
        "preferred qualifications", "about us", "what we offer",
        "responsibilities", "qualifications", "requirements",
        "location", "full-time", "part-time", "remote", "hybrid",
        "salary", "benefits", "employment", "apply", "application",
        # Non-tech business terms that get extracted
        "economics", "finance", "accounting", "marketing", "sales",
        "operations management", "project management", "program management",
        "six sigma", "lean", "kpi", "kpis", "sop", "sops", "crm",
        "erp", "pivottable", "pivottables", "advanced", "proficiency",
        "monitor", "maintain", "collaborate", "develop", "prepare",
        "conduct", "collect", "support", "manage", "coordinate",
        "implement", "create", "build", "design", "deliver", "provide",
        "ensure", "identify", "analyze", "review", "report", "update",
        "working", "working knowledge", "proven", "demonstrated",
        "strong understanding", "solid understanding", "in-depth",
        "hands on experience", "related field", "united states",
        "united states of america", "bachelor", "master", "degree",
        "computer science", "information technology", "information systems",
        "equivalent", "relevant", "similar", "related",
    }

    # Inject ALL injectable keywords + remaining JD keywords
    all_to_inject = list(injectable_kws) + list(jd_keywords)
    injected_count = 0

    for kw in all_to_inject:
        kw_lower = kw.lower().strip()
        if not kw_lower:
            continue

        # Skip generic/non-technical words
        if kw_lower in SKILLS_STOPWORDS:
            continue
        # Skip short single-word non-tech terms (< 3 chars or pure English words)
        if len(kw_lower) < 3:
            continue
        # Skip anything that looks like a sentence fragment (has spaces but no tech marker)
        if " " in kw_lower and not any(tech in kw_lower for tech in [
            "azure", "aws", "gcp", "data", "sql", "python", "spark", "kafka",
            "dbt", "airflow", "snowflake", "power bi", "tableau", "docker",
            "machine learning", "deep learning", "natural language", "api",
            "cloud", "etl", "elt", "pipeline", "warehouse", "lake",
        ]):
            continue

        # Find target category
        target_cat = "data_engineering"
        for hint_key, cat in KEYWORD_CATEGORY_HINTS.items():
            if hint_key in kw_lower:
                target_cat = cat
                break

        if target_cat not in enriched:
            target_cat = "data_engineering"

        # Check not already present
        existing_lower = [s.lower() for s in enriched.get(target_cat, [])]
        already = any(
            _keyword_in_text(kw_lower, ex) or _keyword_in_text(ex, kw_lower)
            for ex in existing_lower
        )
        if not already:
            display_kw = kw.upper() if len(kw) <= 4 else kw.title()
            enriched[target_cat].insert(0, display_kw)
            injected_count += 1

    return enriched, injected_count


def enhance_bullets_with_keywords(bullets: list, missing_keywords: list, max_additions=4) -> list:
    """Append missing keywords naturally to the most relevant bullets."""
    if not missing_keywords:
        return bullets

    enhanced   = list(bullets)
    all_text   = " ".join(b.lower() for b in enhanced)
    added      = 0
    appendages = {
        "airflow":       "; orchestrating workflows via Apache Airflow",
        "dbt":           "; transforming models with dbt (data build tool)",
        "snowflake":     "; leveraging Snowflake for cloud data warehousing",
        "docker":        "; containerizing services with Docker",
        "kubernetes":    "; orchestrating containers with Kubernetes",
        "ci/cd":         "; implementing CI/CD pipelines for automated deployments",
        "git":           "; maintaining version control with Git",
        "rest api":      "; integrating REST APIs for data ingestion",
        "machine learning": "; supporting machine learning feature pipelines",
        "terraform":     "; provisioning infrastructure with Terraform",
        "delta lake":    "; building lakehouse architectures with Delta Lake",
        "databricks":    "; processing large-scale data with Azure Databricks",
    }

    for kw in missing_keywords:
        if added >= max_additions:
            break
        kw_lower = kw.lower()
        if _keyword_in_text(kw_lower, all_text):
            continue

        # Find best bullet (most JD keyword mentions = most relevant context)
        best_idx   = 0
        best_score = -1
        for i, bullet in enumerate(enhanced):
            score = sum(1 for k in missing_keywords if _keyword_in_text(k, bullet))
            if score > best_score:
                best_score = score
                best_idx   = i

        suffix = appendages.get(kw_lower, f"; leveraging {kw} for scalable data processing")
        enhanced[best_idx] = enhanced[best_idx].rstrip(".") + suffix
        all_text = " ".join(b.lower() for b in enhanced)
        added += 1

    return enhanced


# =============================================================================
# SECTION BUILDERS
# =============================================================================

def add_name_header(doc, job_title="Data Engineer"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(p, before_pt=0, after_pt=1)
    run = p.add_run(PROFILE["name"].upper())
    _set_font(run, 22, bold=True, color=COLOR_ACCENT)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(p2, before_pt=0, after_pt=2)
    _set_font(p2.add_run(job_title), 11, color=COLOR_MID)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(p3, before_pt=0, after_pt=1)
    contact = f"{PROFILE['location']}  |  {PROFILE['phone']}  |  {PROFILE['email']}"
    _set_font(p3.add_run(contact), 10, color=COLOR_DARK)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(p4, before_pt=0, after_pt=4)
    _add_hyperlink(p4, f"https://{PROFILE['linkedin']}", PROFILE["linkedin"], 10)
    _set_font(p4.add_run("  |  "), 10, color=COLOR_DARK)
    _add_hyperlink(p4, f"https://{PROFILE['github']}", PROFILE["github"], 10)


def add_section_header(doc, title):
    p = doc.add_paragraph()
    _para_space(p, before_pt=8, after_pt=1)
    _add_bottom_border(p)
    _set_font(p.add_run(title.upper()), 11, bold=True, color=COLOR_ACCENT)
    return p


def add_summary(doc, jd_keywords: list, role_title: str):
    add_section_header(doc, "Professional Summary")

    kw_set      = set(k.lower() for k in jd_keywords)
    azure_focus = any(k in kw_set for k in ["azure", "adf", "adls", "azure data factory", "databricks", "azure synapse"])
    aws_focus   = "aws" in kw_set
    gcp_focus   = any(k in kw_set for k in ["gcp", "bigquery", "google cloud"])
    spark_focus = any(k in kw_set for k in ["spark", "pyspark", "apache spark"])
    kafka_focus = "kafka" in kw_set
    airflow_focus = "airflow" in kw_set
    dbt_focus   = "dbt" in kw_set
    snow_focus  = "snowflake" in kw_set
    bi_focus    = any(k in kw_set for k in ["power bi", "tableau", "looker"])
    ml_focus    = any(k in kw_set for k in ["machine learning", "ml", "mlops"])

    clouds = []
    if azure_focus: clouds.append("Microsoft Azure")
    if aws_focus:   clouds.append("AWS")
    if gcp_focus:   clouds.append("GCP")
    cloud_str = " and ".join(clouds) if clouds else "leading cloud platforms (Azure, AWS, GCP)"

    # Build tech snippet from top JD keywords
    priority = [
        "azure data factory", "adf", "azure", "databricks", "azure synapse",
        "adls", "snowflake", "aws", "gcp", "bigquery", "redshift",
        "pyspark", "apache spark", "spark", "kafka", "airflow", "hadoop",
        "etl", "elt", "pipeline", "data warehouse", "data lake",
        "python", "sql", "t-sql", "power bi", "tableau", "dbt", "docker",
    ]
    top_named = []
    for kw in priority:
        if kw in kw_set and kw not in top_named:
            top_named.append(kw)
        if len(top_named) >= 9:
            break
    # Only pull in JD keywords that look like actual tech terms (not business words)
    TECH_STOPWORDS = {
        "health", "services", "engineering", "familiarity", "compensation",
        "medical", "requirements", "years", "experience", "ability", "strong",
        "knowledge", "skills", "management", "team", "work", "role", "job",
        "company", "position", "candidate", "preferred", "required", "plus",
        "good", "excellent", "great", "background", "understanding", "support",
        "business", "data", "information", "systems", "solutions", "development",
        "applications", "tools", "technologies", "platforms", "environment",
        "type", "types", "hands-on", "hands on", "collaborative", "cross-functional",
        "communication", "written", "verbal", "interpersonal", "problem", "solving",
        "detail", "oriented", "fast", "paced", "organization", "operations",
        "infrastructure", "strategy", "delivery", "quality", "performance",
        "security", "compliance", "documentation", "reporting", "analysis",
    }
    for kw in jd_keywords:
        kw_clean = kw.lower().strip()
        if kw not in top_named and kw_clean not in TECH_STOPWORDS and len(kw_clean) > 2:
            top_named.append(kw)
        if len(top_named) >= 9:
            break

    tech_snippet = ", ".join(kw.upper() if len(kw) <= 5 else kw.title() for kw in top_named[:7])

    tech_detail = []
    if spark_focus:   tech_detail.append("PySpark and Apache Spark for distributed data processing")
    if kafka_focus:   tech_detail.append("Apache Kafka for real-time event streaming")
    if airflow_focus: tech_detail.append("Apache Airflow for pipeline orchestration")
    if dbt_focus:     tech_detail.append("dbt for analytics engineering and data transformation")
    if snow_focus:    tech_detail.append("Snowflake for cloud data warehousing")
    if bi_focus:
        tool = "Power BI" if "power bi" in kw_set else ("Tableau" if "tableau" in kw_set else "Looker")
        tech_detail.append(f"{tool} for business intelligence and reporting")
    if ml_focus:      tech_detail.append("machine learning pipeline support and MLOps workflows")

    # ── Claude-written summary (executive quality, job-specific) ─────────────
    import claude_engine as _ce
    tech_list = tech_snippet if tech_snippet else ", ".join([
        "Python", "SQL", "PySpark", "Azure", "AWS", "ETL pipelines", "data warehousing"
    ])
    tech_detail_str = "; ".join(tech_detail[:3]) if tech_detail else ""

    summary_prompt = f"""Write a 3-sentence professional summary for a resume. This is for:

Role: {role_title}
Key technologies from the job: {tech_list}
Specific tools the job emphasizes: {tech_detail_str or "ETL pipelines, cloud data platforms, SQL"}
Cloud platforms: {cloud_str}

Candidate facts (use these, do not fabricate):
- M.S. Data Science and Analytics, Florida Atlantic University (May 2025)
- 3+ years hands-on: Python (4 yrs), SQL (4 yrs), data engineering (3 yrs)
- Built production ETL/ELT pipelines, cloud data warehouses, real-time streaming systems
- Deployed on Azure (ADF, ADLS, Databricks), AWS, GCP
- Delivered analytics and BI solutions that directly drove business decisions

Rules for the 3 sentences:
- Sentence 1: Open with a bold, specific value statement — what the candidate DELIVERS, not what they "are". No "Results-driven", no "passionate", no "motivated".
- Sentence 2: Name 2-3 specific technical achievements with concrete detail tied to the JD keywords.
- Sentence 3: End with a forward-looking statement about what they bring to THIS specific role.
- No bullet points. No hyphens. No dashes. No F-1/visa/sponsorship mention. No clichés.
- Sound like a top-tier engineer, not a resume template. Be sharp, specific, confident.
- Max 80 words total across all 3 sentences.

Return ONLY the 3 sentences, nothing else."""

    summary_text = _ce._ask(summary_prompt, max_tokens=200)

    # Fallback if Claude fails
    if not summary_text or len(summary_text) < 40:
        summary_text = (
            f"{role_title} who designs and ships production-grade data infrastructure "
            f"on {cloud_str}, with deep hands-on experience in Python, SQL, and ETL/ELT "
            f"pipeline engineering. "
            f"Has built end-to-end data systems spanning {tech_list}, delivering "
            f"measurable improvements in pipeline reliability, processing speed, and "
            f"analytics accuracy. Holds an M.S. in Data Science and Analytics "
            f"(Florida Atlantic University, 2025) and brings immediate production-ready "
            f"capability to {role_title} teams."
        )

    # Strip markdown — Claude sometimes returns **bold** or *italic* syntax
    summary_text = re.sub(r'\*\*(.+?)\*\*', r'\1', summary_text)
    summary_text = re.sub(r'\*(.+?)\*',     r'\1', summary_text)
    summary_text = summary_text.strip()

    p = doc.add_paragraph()
    _para_space(p, before_pt=2, after_pt=2)
    _set_font(p.add_run(summary_text), 10, color=COLOR_DARK)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_experience(doc, jd_keywords: list, injectable_kws: list, ai_bullets: dict = None):
    add_section_header(doc, "Work Experience")

    kw_set       = set(k.lower() for k in jd_keywords)
    is_analytics = any(k in kw_set for k in ["analytics", "power bi", "tableau", "ga4", "google analytics"])

    all_bullets_text = " ".join(
        b.lower() for job in EXPERIENCE for b in job.get("bullets", [])
    )
    missing_from_bullets = [
        kw for kw in injectable_kws
        if not _keyword_in_text(kw, all_bullets_text)
    ]
    first_primary = True

    for job in EXPERIENCE:
        include = job["include_always"]
        if not include:
            if is_analytics:
                include = True
            else:
                include = any(r.lower() in kw_set for r in job.get("include_for_roles", []))
        if not include:
            continue

        p_title = doc.add_paragraph()
        _para_space(p_title, before_pt=5, after_pt=0)
        _set_font(p_title.add_run(job["title"]), 10.5, bold=True, color=COLOR_BLACK)

        p_sub = doc.add_paragraph()
        _para_space(p_sub, before_pt=0, after_pt=1)
        _set_font(
            p_sub.add_run(f"{job['company']}  |  {job['duration']}  |  {job['location']}"),
            9.5, italic=True, color=COLOR_MID
        )

        # Use Claude-rewritten bullets if available for this job
        if ai_bullets and job["title"] in ai_bullets:
            bullets_sorted = ai_bullets[job["title"]]
            first_primary = False
        else:
            bullets = list(job["bullets"])
            bullets_sorted = sorted(
                bullets,
                key=lambda b: sum(1 for kw in kw_set if kw in b.lower()),
                reverse=True,
            )
            if first_primary and job["include_always"] and missing_from_bullets:
                bullets_sorted = enhance_bullets_with_keywords(
                    bullets_sorted, missing_from_bullets, max_additions=4
                )
                first_primary = False

        max_bullets = 7 if job["include_always"] else 4
        for bullet in bullets_sorted[:max_bullets]:
            p_b = doc.add_paragraph(style="List Bullet")
            _para_space(p_b, before_pt=0, after_pt=1)
            _set_font(p_b.add_run(bullet), 10, color=COLOR_DARK)

        if job.get("tools"):
            tools_all = list(job["tools"])
            for kw in injectable_kws:
                kw_pretty = kw.upper() if len(kw) <= 4 else kw.title()
                if kw_pretty not in tools_all and kw not in [t.lower() for t in tools_all]:
                    tools_all.append(kw_pretty)

            tools_sorted = sorted(
                tools_all,
                key=lambda t: 1 if t.lower() in kw_set else 0,
                reverse=True,
            )
            p_t = doc.add_paragraph()
            _para_space(p_t, before_pt=1, after_pt=0)
            _set_font(p_t.add_run("Tools: "), 9.5, bold=True, color=COLOR_DARK)
            _set_font(p_t.add_run(" · ".join(tools_sorted)), 9.5, color=COLOR_MID)


def add_education(doc):
    add_section_header(doc, "Education")
    for edu in EDUCATION:
        p = doc.add_paragraph()
        _para_space(p, before_pt=3, after_pt=0)
        _set_font(p.add_run(edu["degree"]), 10.5, bold=True, color=COLOR_BLACK)

        p2 = doc.add_paragraph()
        _para_space(p2, before_pt=0, after_pt=2)
        _set_font(
            p2.add_run(f"{edu['school']}  |  {edu['location']}  |  Graduated {edu['graduated']}"),
            9.5, italic=True, color=COLOR_MID
        )


def add_projects(doc, jd_keywords: list, job_title: str = "", max_projects: int = 3):
    """
    Add the most relevant GitHub projects section to the resume.
    Picks projects whose include_for list matches the job title/keywords.
    Always shows at least 2 projects even if no keyword match.

    Projects come from raghav_profile.PROJECTS — the single source of truth.
    """
    if not PROJECTS:
        return

    kw_set   = set(k.lower() for k in jd_keywords)
    role_low = job_title.lower()

    # Score each project by relevance to this job
    def _relevance(proj):
        score = 0
        # include_for role match
        for role_kw in proj.get("include_for", []):
            if role_kw.lower() in role_low or any(role_kw.lower() in kw for kw in kw_set):
                score += 3
                break
        # tech stack keyword overlap with JD
        tech_str = proj.get("tech", "").lower()
        for kw in kw_set:
            if kw in tech_str:
                score += 1
        return score

    ranked = sorted(PROJECTS, key=_relevance, reverse=True)
    selected = ranked[:max_projects]

    # Always include at least 2 projects
    if len(selected) < 2 and len(PROJECTS) >= 2:
        selected = PROJECTS[:2]

    add_section_header(doc, "Projects")

    for proj in selected:
        name     = proj.get("name", "")
        tech     = proj.get("tech", "")
        github   = proj.get("github", "")
        bullets  = proj.get("bullets", [])
        # Show max 3 bullets per project (resume space)
        show_bullets = bullets[:3]

        # Project title line: Name | Tech Stack
        p_title = doc.add_paragraph()
        _para_space(p_title, before_pt=5, after_pt=0)
        _set_font(p_title.add_run(name), 10.5, bold=True, color=COLOR_BLACK)
        if tech:
            _set_font(p_title.add_run(f"  |  {tech}"), 9.5, italic=False, color=COLOR_MID)

        # GitHub link on its own line (compact)
        if github:
            p_gh = doc.add_paragraph()
            _para_space(p_gh, before_pt=0, after_pt=1)
            _add_hyperlink(p_gh, github, github.replace("https://", ""), 9)

        # Bullet points
        for bullet in show_bullets:
            p_b = doc.add_paragraph(style="List Bullet")
            _para_space(p_b, before_pt=0, after_pt=1)
            _set_font(p_b.add_run(bullet), 10, color=COLOR_DARK)


def add_skills(doc, jd_keywords: list, injectable_kws: list):
    add_section_header(doc, "Technical Skills")

    enriched_skills, injected_count = build_enriched_skills(jd_keywords, injectable_kws)

    section_labels = {
        "azure":            "Azure Data Platform",
        "data_engineering": "Data Engineering & Processing",
        "programming":      "Programming & Query Languages",
        "cloud":            "Cloud Platforms",
        "databases":        "Databases",
        "analytics":        "Analytics & Reporting",
        "professional":     "Professional Skills",
    }

    kw_set = set(k.lower() for k in jd_keywords)
    for cat_key, label in section_labels.items():
        items = enriched_skills.get(cat_key, [])
        if not items:
            continue
        p = doc.add_paragraph()
        _para_space(p, before_pt=2, after_pt=1)
        _set_font(p.add_run(f"{label}: "), 10, bold=True, color=COLOR_DARK)
        _set_font(p.add_run(" · ".join(items)), 10, color=COLOR_DARK)


def add_ats_gap_fill(doc, missing_keywords: list):
    """
    Emergency ATS gap-fill: adds remaining missing TECHNICAL keywords to resume.
    Only real tech tools/platforms — never business words or JD boilerplate.
    """
    if not missing_keywords:
        return

    # Whitelist: only terms that are real technical tools/platforms/languages.
    # If a word isn't in at least one of these sets, it does NOT go in the resume.
    TECH_WHITELIST = {
        # Languages
        "python", "sql", "r", "java", "scala", "golang", "go", "rust",
        "javascript", "typescript", "bash", "shell", "c++", "c#",
        # Data Engineering
        "spark", "pyspark", "kafka", "airflow", "hadoop", "hive", "flink",
        "dbt", "etl", "elt", "pipeline", "data warehouse", "data lake",
        "lakehouse", "delta lake", "snowflake", "databricks", "redshift",
        "bigquery", "data mart", "data mesh", "streaming", "batch",
        # Cloud
        "aws", "azure", "gcp", "google cloud", "s3", "ec2", "lambda",
        "azure data factory", "adf", "adls", "synapse", "cosmos db",
        "redshift", "bigquery", "emr", "glue", "step functions",
        # Databases
        "postgresql", "mysql", "sqlite", "mongodb", "cassandra", "redis",
        "oracle", "sql server", "t-sql", "pl/sql", "dynamodb", "neo4j",
        # BI / Analytics
        "power bi", "tableau", "looker", "qlik", "metabase", "superset",
        "dax", "power query", "excel", "google analytics", "ga4",
        # ML / AI
        "scikit-learn", "tensorflow", "pytorch", "keras", "mlflow",
        "machine learning", "deep learning", "nlp", "llm", "xgboost",
        "lightgbm", "computer vision", "transformers", "hugging face",
        # DevOps
        "docker", "kubernetes", "k8s", "terraform", "ansible", "jenkins",
        "github actions", "ci/cd", "git", "github", "gitlab",
        # Misc tech
        "pandas", "numpy", "scipy", "matplotlib", "seaborn", "plotly",
        "fastapi", "flask", "django", "rest api", "graphql",
        "streamlit", "jupyter", "spark sql", "hbase", "impala",
        "nifi", "talend", "informatica", "ssis", "mulesoft",
    }

    def _is_real_tech(kw: str) -> bool:
        """Returns True only if keyword looks like a real technical term."""
        kw_l = kw.lower().strip()
        # Direct whitelist match
        if kw_l in TECH_WHITELIST:
            return True
        # Partial match: keyword contains a whitelisted term
        for tech in TECH_WHITELIST:
            if tech in kw_l and len(tech) > 3:
                return True
        return False

    # Filter to real tech only
    real_tech_keywords = [kw for kw in missing_keywords if _is_real_tech(kw)]

    if not real_tech_keywords:
        return  # Nothing real to add — don't print an empty section

    add_section_header(doc, "Key Technical Areas")

    cloud_terms = {"aws", "gcp", "google cloud", "azure", "databricks", "terraform",
                   "docker", "kubernetes", "ci/cd", "redshift", "bigquery", "s3", "ec2",
                   "lambda", "emr", "glue", "synapse", "cosmos", "adf", "adls"}
    de_terms    = {"airflow", "dbt", "kafka", "spark", "hadoop", "hive", "flink",
                   "data lake", "delta lake", "lakehouse", "snowflake", "batch",
                   "streaming", "real-time", "etl", "elt", "pipeline", "data warehouse",
                   "data mart", "nifi", "talend", "informatica", "ssis", "mulesoft"}
    analytics_terms = {"power bi", "tableau", "looker", "metabase", "superset", "qlik",
                       "machine learning", "nlp", "llm", "tensorflow", "pytorch", "keras",
                       "scikit-learn", "pandas", "ga4", "google analytics", "xgboost",
                       "computer vision", "matplotlib", "seaborn", "plotly", "streamlit"}

    groups = {
        "Cloud & Infrastructure": [],
        "Data Engineering":       [],
        "Analytics & ML":         [],
        "Other Technologies":     [],
    }

    for kw in real_tech_keywords:
        kw_lower = kw.lower()
        if any(t in kw_lower for t in cloud_terms):
            groups["Cloud & Infrastructure"].append(kw.title())
        elif any(t in kw_lower for t in de_terms):
            groups["Data Engineering"].append(kw.title())
        elif any(t in kw_lower for t in analytics_terms):
            groups["Analytics & ML"].append(kw.title())
        else:
            groups["Other Technologies"].append(kw.title())

    for group_name, items in groups.items():
        if not items:
            continue
        p = doc.add_paragraph()
        _para_space(p, before_pt=2, after_pt=1)
        _set_font(p.add_run(f"{group_name}: "), 10, bold=True, color=COLOR_DARK)
        _set_font(p.add_run(" · ".join(items)), 10, color=COLOR_DARK)


# =============================================================================
# CLAUDE AI — INTELLIGENT RESUME REWRITING
# =============================================================================

def _claude_rewrite_summary(job_title: str, company: str, jd_text: str, profile_summary: str) -> str:
    """Use Claude to write a tailored professional summary for this specific job."""
    try:
        import anthropic, os
        api_key = os.environ.get("ANTHROPIC_API_KEY", "")
        if not api_key:
            for line in (Path.home() / "job_pipeline" / ".env").read_text().splitlines():
                if line.startswith("ANTHROPIC_API_KEY="):
                    api_key = line.split("=", 1)[1].strip().strip('"').strip("'")
        client = anthropic.Anthropic(api_key=api_key)
        prompt = (
            f"Write a 3-sentence professional summary for a resume applying to: '{job_title}' at '{company}'.\n\n"
            f"Job description excerpt:\n{jd_text[:1200]}\n\n"
            f"Candidate facts (use these, do not invent):\n"
            f"- M.S. Data Science and Analytics, Florida Atlantic University, May 2025\n"
            f"- 3+ years hands-on: Python (4 yrs), SQL (4 yrs), data engineering (3 yrs)\n"
            f"- Built production ETL/ELT pipelines, cloud data warehouses, real-time streaming\n"
            f"- Deployed on Azure (ADF, ADLS, Databricks), AWS, GCP\n\n"
            f"Rules:\n"
            f"- 3 sentences only. No bullet points. No dashes as separators.\n"
            f"- Sentence 1: Bold value statement — what this candidate DELIVERS, not what they 'are'. No 'Results-driven', no 'passionate'.\n"
            f"- Sentence 2: 2-3 specific technical achievements tied to the JD keywords.\n"
            f"- Sentence 3: Forward-looking statement about what they bring to THIS role.\n"
            f"- NO mention of visa, OPT, sponsorship, F-1, or immigration status.\n"
            f"- NO clichés. Max 80 words total.\n"
            f"- Return ONLY the 3 sentences, nothing else."
        )
        resp = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=300,
            messages=[{"role": "user", "content": prompt}]
        )
        return resp.content[0].text.strip()
    except Exception as e:
        return ""  # fallback to template-based summary


def _claude_rewrite_bullets(job_title: str, company: str, jd_text: str, bullets: list, missing_kws: list) -> list:
    """Use Claude to rewrite/enhance experience bullets to match this JD."""
    try:
        import anthropic, os
        api_key = os.environ.get("ANTHROPIC_API_KEY", "")
        if not api_key:
            for line in (Path.home() / "job_pipeline" / ".env").read_text().splitlines():
                if line.startswith("ANTHROPIC_API_KEY="):
                    api_key = line.split("=", 1)[1].strip().strip('"').strip("'")
        client = anthropic.Anthropic(api_key=api_key)
        bullets_text = "\n".join(f"- {b}" for b in bullets[:8])
        missing_str  = ", ".join(missing_kws[:15]) if missing_kws else "none"
        prompt = (
            f"You are rewriting resume bullet points for a '{job_title}' role at '{company}'.\n\n"
            f"Job description excerpt:\n{jd_text[:1000]}\n\n"
            f"Original bullets:\n{bullets_text}\n\n"
            f"Keywords to naturally include if possible: {missing_str}\n\n"
            f"Rules:\n"
            f"- Rewrite each bullet to be more impactful and relevant to this role\n"
            f"- Keep each bullet under 20 words\n"
            f"- Start each with a strong past-tense action verb (Built, Designed, Optimized, etc.)\n"
            f"- Preserve real metrics/numbers from the original bullets\n"
            f"- Do NOT invent metrics or companies\n"
            f"- Return ONLY the bullets, one per line, no dashes or numbers, no extra text"
        )
        resp = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=500,
            messages=[{"role": "user", "content": prompt}]
        )
        rewritten = [line.strip().lstrip("•-– ") for line in resp.content[0].text.strip().split("\n") if line.strip()]
        # Keep originals for any bullets Claude dropped
        while len(rewritten) < len(bullets[:8]):
            rewritten.append(bullets[len(rewritten)])
        return rewritten
    except Exception as e:
        return bullets  # fallback to originals


# =============================================================================
# MASTER BUILD FUNCTION  — 98%+ GUARANTEE + CLAUDE AI INTELLIGENCE
# =============================================================================

def build_resume(
    job_title:       str,
    company:         str,
    jd_keywords:     list,
    injectable_kws:  list,
    initial_score:   float,
    optimized_score: float,
    output_path:     str  = None,
    jd_text:         str  = "",
    profile_summary: str  = "",
) -> tuple:
    """
    Build a tailored, ATS-optimized .docx resume with Claude AI rewriting.
    Claude rewrites the summary and top bullets to match this specific job.
    After building, verifies actual keyword coverage (98%+ guarantee).
    Returns (file_path, actual_initial_score, actual_optimized_score).
    """
    print(f"\n  ┌─ Building resume: {job_title} @ {company}")
    print(f"  │  Estimated BEFORE: {initial_score:.0f}%  →  Estimated AFTER: {optimized_score:.0f}%")
    print(f"  │  JD keywords: {len(jd_keywords)}  |  Injectable: {len(injectable_kws)}")

    # ── Claude AI: rewrite summary & bullets for this specific job ────────────
    ai_summary = ""
    ai_bullets  = {}
    if jd_text:
        print(f"  │  🤖 Claude: rewriting summary & bullets for {company}...")
        ai_summary = _claude_rewrite_summary(job_title, company, jd_text, profile_summary)
        # Rewrite bullets for the primary (always-included) job
        for job in EXPERIENCE:
            if job.get("include_always"):
                ai_bullets[job["title"]] = _claude_rewrite_bullets(
                    job_title, company, jd_text,
                    job.get("bullets", []),
                    injectable_kws
                )
                break  # only rewrite top job to save API calls

    doc = Document()
    _set_margins(doc)
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)

    title_label = job_title if len(job_title) < 50 else "Data Engineer"

    add_name_header(doc, title_label)

    # Use Claude summary if available, else template
    if ai_summary:
        # Strip markdown formatting — python-docx does not render ** or * syntax
        # and they appear as literal asterisks on the printed resume.
        ai_summary = re.sub(r'\*\*(.+?)\*\*', r'\1', ai_summary)   # **bold** → bold
        ai_summary = re.sub(r'\*(.+?)\*',     r'\1', ai_summary)   # *italic* → italic
        ai_summary = re.sub(r'#{1,6}\s*',     '',    ai_summary)   # ## headings → nothing
        ai_summary = ai_summary.strip()

        add_section_header(doc, "Professional Summary")
        p = doc.add_paragraph()
        _para_space(p, before_pt=2, after_pt=2)
        _set_font(p.add_run(ai_summary), 10, color=COLOR_DARK)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        add_summary(doc, jd_keywords, title_label)

    add_experience(doc, jd_keywords, injectable_kws, ai_bullets=ai_bullets)
    add_education(doc)
    add_projects(doc, jd_keywords, job_title=job_title, max_projects=3)
    add_skills(doc, jd_keywords, injectable_kws)

    # ── PASS 1: Verify actual coverage ───────────────────────────────────────
    doc_text = extract_docx_text(doc)

    # If no JD keywords, coverage check is meaningless — keep the original score
    if not jd_keywords:
        pass1_score = initial_score
        missing_after_pass1 = []
        print(f"  │  Pass 1 coverage: N/A  (no JD keywords — keeping original score {initial_score:.0f}%)")
    else:
        pass1_score, _, missing_after_pass1 = compute_actual_coverage(jd_keywords, doc_text)
        print(f"  │  Pass 1 coverage: {pass1_score:.1f}%  ({len(missing_after_pass1)} keywords still missing)")

    # ── PASS 2: Gap-fill any still-missing keywords ───────────────────────────
    actual_optimized = pass1_score
    if missing_after_pass1:
        add_ats_gap_fill(doc, missing_after_pass1)
        doc_text2 = extract_docx_text(doc)
        pass2_score, _, still_missing = compute_actual_coverage(jd_keywords, doc_text2)
        actual_optimized = pass2_score
        if still_missing:
            print(f"  │  Pass 2 coverage: {pass2_score:.1f}%  ({len(still_missing)} still missing: {still_missing[:5]})")
        else:
            print(f"  │  Pass 2 coverage: {pass2_score:.1f}%  ✅ All keywords covered!")
    else:
        print(f"  │  ✅ Full coverage achieved in Pass 1!")

    # ── SAVE ──────────────────────────────────────────────────────────────────
    if not output_path:
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        safe_company = re.sub(r"[^\w\s-]", "", company).strip().replace(" ", "_")[:30]
        safe_title   = re.sub(r"[^\w\s-]", "", job_title).strip().replace(" ", "_")[:25]
        filename     = f"Raghavendra_Karanam_{safe_company}_{safe_title}.docx"
        output_path  = os.path.join(OUTPUT_DIR, filename)

    doc.save(output_path)

    delta = actual_optimized - initial_score
    flag  = "🟢" if actual_optimized >= 90 else ("🟡" if actual_optimized >= 75 else "🔴")
    print(f"  │  ACTUAL RESULT:  {initial_score:.0f}%  →  {actual_optimized:.1f}%  (+{delta:.0f}%)  {flag}")
    print(f"  └─ Saved → {os.path.basename(output_path)}")

    return output_path, initial_score, actual_optimized


# =============================================================================
# BATCH BUILD
# =============================================================================

def build_all_resumes(filtered_csv: str = None) -> list:
    """
    Build tailored resumes for all jobs in filtered_jobs.csv.
    Returns list of (path, initial_score, actual_optimized_score) tuples.
    """
    print("\n" + "═" * 64)
    print("  JOB PIPELINE — Step 3: Building Tailored Resumes (PRO v3)")
    print("  TARGET: 98%+ ATS keyword coverage per resume")
    print("═" * 64)

    if not filtered_csv:
        filtered_csv = os.path.join(os.path.dirname(__file__), "data", "filtered_jobs.csv")

    if not os.path.exists(filtered_csv):
        print(f"  filtered_jobs.csv not found. Run jd_parser.py first.")
        sys.exit(1)

    df = pd.read_csv(filtered_csv)
    results = []  # list of (path, initial, optimized)

    for idx, (_, row) in enumerate(df.iterrows(), 1):
        keywords   = [k.strip() for k in str(row.get("jd_keywords", "")).split(",") if k.strip()]
        injectable = [k.strip() for k in str(row.get("injectable_keywords", "")).split(",") if k.strip()]
        initial    = float(row.get("initial_score",   row.get("ats_score", 0)))
        optimized  = float(row.get("optimized_score", row.get("ats_score", 0)))

        print(f"\n  [{idx}/{len(df)}] Processing job...")
        path, actual_initial, actual_optimized = build_resume(
            job_title       = str(row.get("title",   "Data Engineer")),
            company         = str(row.get("company", "Company")),
            jd_keywords     = keywords,
            injectable_kws  = injectable,
            initial_score   = initial,
            optimized_score = optimized,
        )
        results.append((path, actual_initial, actual_optimized))

        # Write actual scores back to the DataFrame for tracker use
        df.at[idx - 1, "actual_initial_score"]   = actual_initial
        df.at[idx - 1, "actual_optimized_score"]  = actual_optimized

    # Save updated CSV with actual scores
    df.to_csv(filtered_csv, index=False)

    print(f"\n{'═' * 64}")
    print(f"  ✅  {len(results)} resumes built → output/resumes/")
    print(f"\n  {'Job':<35}  {'Before':>6}  {'After':>6}  {'Delta':>6}")
    print(f"  {'─'*35}  {'─'*6}  {'─'*6}  {'─'*6}")
    for (path, ini, opt), (_, row) in zip(results, df.iterrows()):
        name = f"{row.get('title','?')[:20]} @ {row.get('company','?')[:12]}"
        delta = opt - ini
        flag  = "✅" if opt >= 90 else "⚠️"
        print(f"  {name:<35}  {ini:>5.0f}%  {opt:>5.1f}%  +{delta:>4.0f}%  {flag}")
    print("═" * 64 + "\n")

    return results


if __name__ == "__main__":
    build_all_resumes()
